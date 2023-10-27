import re, docx
import sys
import os
import docx
from docx import Document
from docx.shared import Inches

def insert_text_in_table(doc, paragraph_index, text):
    # Add a table with one row and one column after the target paragraph
    table = doc.add_table(rows=1, cols=1)
    
    # Move the table below the target paragraph
    for _ in range(len(doc.paragraphs) - 2, paragraph_index, -1):
        move_table_after(doc.paragraphs[_], table)
    
    # No border for aesthetics
    set_table_border(table, False)
    
    # Inserting text
    cell = table.cell(0, 0)
    cell.text = text

def set_table_border(table, has_border):
    # A function to set or remove borders from a table
    for row in table.rows:
        for cell in row.cells:
            for key, value in cell._element.attrib.items():
                if 'border' in key:
                    cell._element.attrib[key] = '1' if has_border else '0'

def move_table_after(paragraph, table):
    """
    Move the given table after the given paragraph.
    """
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def find_date_paragraph_index(doc):
    """
    This function finds and returns the index of the paragraph containing the date.
    """
    date_pattern = re.compile(r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}\b')
    for i, para in enumerate(doc.paragraphs):
        if date_pattern.search(para.text):
            return i  # Return the index instead of the paragraph object
    return None

def mark_revision(doc, keyword="Revised"):
    """
    This function checks the paragraph following the date for the revision keyword 
    and updates the text based on whether it is already revised.
    """
    date_para_index = find_date_paragraph_index(doc)

    # Initialize to 1, assuming it's the first revision if no previous keyword is found
    revision_number = 1  

    # If there's no date paragraph, we'll still proceed with the assumption that it's a valid document for revision.
    if date_para_index is not None and date_para_index + 1 < len(doc.paragraphs):
        target_para = doc.paragraphs[date_para_index + 1]
        text = target_para.text.split()

        for i, word in enumerate(text):
            if word == keyword and i + 1 < len(text) and text[i + 1].isdigit():
                revision_number = int(text[i + 1]) + 1  # increment the revision number
                text[i + 1] = str(revision_number)
                break
        else:
            # If "Revised" keyword is present but without a number, add the number.
            if keyword in target_para.text:
                text.append(str(revision_number))
            else:
                # If the "Revised" keyword doesn't exist, add "Revised 1" at the end.
                text.append(f"{keyword} 1")

        # Update paragraph text with the tabs
        tabs = '\t' * 6
        new_text = ' '.join(text)
        target_para.clear()
        target_para.add_run(tabs + new_text)  # Adding the new text with tabs

    else:
        # If we don't find a paragraph following the date, we will create a new paragraph with the revision note.
        tabs = '\t' * 6
        new_para = doc.add_paragraph()
        new_para.add_run(tabs + f"{keyword} {revision_number}")  # Adding the new text with tabs

    return revision_number  # Return the revision number

def set_tabs_in_paragraph(paragraph, tab_positions):
    # Clear existing tab stops
    paragraph.paragraph_format.tab_stops.clear_all()
    
    # Add new tab stops at specified positions
    for pos in tab_positions:
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(pos))
# # Usage in your code
# tab_positions = [3, 6]  # positions in inches
# set_tabs_in_paragraph(target_paragraph, tab_positions)
# target_paragraph.add_run('\t').add_run("Your Text Here")
def main(file_path):  # Now main accepts an argument, which is the file path of the document.
    try:
        # Load the existing document
        doc = Document(file_path)
        index = find_date_paragraph_index(doc) + 1 if find_date_paragraph_index(doc) is not None else None

        # Mark the document as revised
        revision_number = mark_revision(doc)
        
        if revision_number is not None:
                # If a revision was made, save the document with the new revision number in the name
                insert_text_in_table(doc, index, f"Revised {revision_number}")
                revised_file_path = os.path.splitext(file_path)[0] + f" Rev {revision_number}" + os.path.splitext(file_path)[1]
                doc.save(revised_file_path)
                print(f"Document revised and saved as '{revised_file_path}'")
                return revised_file_path

    except Exception as e:
        print(f"An error occurred: {e}")
        return None  # Return None or a suitable message indicating the failure.

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Error: No file path provided.")
        sys.exit(1)

    # sys.argv[1] will be the file path argument passed from the command line.
    file_path = sys.argv[1]
    revised_file_path = main(file_path)
    print(revised_file_path)